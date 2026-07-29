#!/usr/bin/env python3
"""Extract text from common requirement/design document formats."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path


def read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="replace")


def read_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - dependency message
        raise RuntimeError("读取 .docx 需要 python-docx") from exc

    document = Document(path)
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))

    return "\n".join(chunks)


def read_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except Exception as exc:  # pragma: no cover - dependency message
        raise RuntimeError("读取 .xlsx 需要 openpyxl") from exc

    workbook = load_workbook(path, read_only=True, data_only=True)
    chunks: list[str] = []

    for sheet in workbook.worksheets:
        chunks.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                chunks.append(" | ".join(values))

    return "\n".join(chunks)


def read_pdf(path: Path) -> str:
    try:
        import pdfplumber  # type: ignore
    except Exception:
        pdfplumber = None

    if pdfplumber is not None:
        chunks: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    chunks.append(f"# Page {page_index}\n{text.strip()}")
        return "\n\n".join(chunks)

    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:
        raise RuntimeError("读取 .pdf 需要安装 pypdf 或 pdfplumber") from exc

    reader = PdfReader(str(path))
    chunks = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            chunks.append(f"# Page {page_index}\n{text.strip()}")
    return "\n\n".join(chunks)


def extract(path: Path) -> dict:
    suffix = path.suffix.lower()
    if not path.exists():
        return {"path": str(path), "ok": False, "error": "文件不存在"}
    if not path.is_file():
        return {"path": str(path), "ok": False, "error": "不是文件"}

    try:
        if suffix in {".md", ".txt"}:
            content = read_text(path)
        elif suffix == ".docx":
            content = read_docx(path)
        elif suffix == ".xlsx":
            content = read_xlsx(path)
        elif suffix == ".pdf":
            content = read_pdf(path)
        else:
            return {"path": str(path), "ok": False, "error": f"不支持的文件类型: {suffix}"}
    except Exception as exc:
        return {"path": str(path), "ok": False, "error": str(exc)}

    return {
        "path": str(path),
        "name": path.name,
        "suffix": suffix,
        "ok": True,
        "chars": len(content),
        "content": content,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract testcase source documents to JSON.")
    parser.add_argument("--input", action="append", required=True, help="Document path. Repeat for multiple files.")
    parser.add_argument("--output", required=True, help="Output JSON path.")
    args = parser.parse_args()

    documents = [extract(Path(item).expanduser()) for item in args.input]
    output = Path(args.output).expanduser()
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps({"documents": documents}, ensure_ascii=False, indent=2), encoding="utf-8")

    failed = [doc for doc in documents if not doc.get("ok")]
    if failed:
        for doc in failed:
            print(f"[WARN] {doc['path']}: {doc['error']}", file=sys.stderr)
    print(str(output))
    return 0 if len(failed) < len(documents) else 1


if __name__ == "__main__":
    raise SystemExit(main())
